"""
data_preprocessor.py

功能：
- 清理 description 字段中的 HTML 样式信息，保留有意义的文字内容、超链接和图片地址
- 使用 DeepSeek API 对内容进行准确复述
- 处理腾讯文档链接和图片内容（预留功能）
- 优化 TAPD 数据的 description 字段

依赖：aiohttp, beautifulsoup4, python-docx
"""
import os
import json
import re
import aiohttp
import asyncio
from typing import Dict, List, Optional, Any
from bs4 import BeautifulSoup
from mcp.server.fastmcp import FastMCP
from docx import Document
import csv
import shutil
import uuid
from .common_utils import get_api_manager, get_file_manager, get_token_counter, TokenBudgetUtils

mcp = FastMCP("data_preprocessor")

async def call_deepseek_api(content: str, session: aiohttp.ClientSession) -> str:
    """调用 DeepSeek/SF 复述接口，并基于总上下文预算动态分配输出tokens。"""
    api_manager = get_api_manager()
    token_counter = get_token_counter()

    prompt = f"""适当简化以下文本，保留关键信息：

{content}

只输出简化结果，不要解释。"""

    # 动态计算输出token预算（统一工具）：总预算=6000
    max_out = TokenBudgetUtils.compute_response_tokens(
        prompt,
        total_budget=6000,
        desired_response_cap=2000,
    )

    return await api_manager.call_llm(prompt, session, max_tokens=max_out)

def clean_html_styles(html_content: str) -> str:
    """清理HTML样式信息，保留有意义的文字内容、超链接和图片地址"""
    if not html_content:
        return ""
    
    # 先用正则表达式清理样式属性
    import re
    
    # 移除 style 属性
    html_content = re.sub(r'\s*style="[^"]*"', '', html_content)
    html_content = re.sub(r"\s*style='[^']*'", '', html_content)
    
    # 移除其他无用属性，但保留 href、src、alt、title
    # 移除 data-* 属性
    html_content = re.sub(r'\s*data-[^=]*="[^"]*"', '', html_content)
    html_content = re.sub(r"\s*data-[^=]*='[^']*'", '', html_content)
    
    # 移除常见的无用属性
    useless_attrs = ['margin', 'padding', 'line-height', 'color', 'font-family', 'font-size', 
                     'font-style', 'font-variant', 'font-weight', 'letter-spacing', 'orphans',
                     'text-align', 'text-indent', 'text-transform', 'widows', 'word-spacing',
                     'webkit-text-stroke-width', 'white-space', 'background-color', 
                     'text-decoration-thickness', 'text-decoration-style', 'text-decoration-color']
    
    for attr in useless_attrs:
        html_content = re.sub(f'\\s*{attr}="[^"]*"', '', html_content)
        html_content = re.sub(f"\\s*{attr}='[^']*'", '', html_content)
    
    # 使用 BeautifulSoup 解析清理后的 HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 提取有用信息
    result_parts = []
    
    # 处理文本内容
    text_content = soup.get_text(separator=' ', strip=True)
    # 清理多余的空格和无用文本
    text_content = re.sub(r'\s+', ' ', text_content)
    text_content = text_content.replace('...', '').strip()
    if text_content:
        result_parts.append(text_content)
    
    # 提取超链接 - 使用正则表达式直接提取
    link_pattern = r'<a[^>]*href=[\'"](https?://[^\'">]+)[\'"][^>]*>([^<]*)</a>'
    links = re.findall(link_pattern, html_content, re.IGNORECASE)
    for href, text in links:
        if text.strip():
            result_parts.append(f"链接: {text.strip()} ({href})")
        else:
            result_parts.append(f"链接: {href}")
    
    # 提取图片信息 - 使用正则表达式直接提取
    img_pattern = r'<img[^>]*src=[\'"](/?[^\'">]+)[\'"][^>]*(?:alt=[\'"]([^\'">]*)[\'"])?[^>]*>'
    images = re.findall(img_pattern, html_content, re.IGNORECASE)
    for src, alt in images:
        if alt:
            result_parts.append(f"图片: {alt} ({src})")
        else:
            result_parts.append(f"图片: {src}")
    
    return ' '.join(result_parts)

def extract_document_links(content: str) -> List[str]:
    """提取腾讯文档链接"""
    # 匹配腾讯文档链接模式
    tencent_doc_pattern = r'https://docs\.qq\.com/[^\s\)"<>]+|https://doc\.weixin\.qq\.com/[^\s\)"<>]+'
    links = re.findall(tencent_doc_pattern, content)
    # 清理链接末尾的引号等字符
    cleaned_links = []
    for link in links:
        # 移除末尾的引号、尖括号等
        link = re.sub(r'["\'>]+$', '', link)
        if link and link not in cleaned_links:
            cleaned_links.append(link)
    return cleaned_links

def extract_image_paths(content: str) -> List[str]:
    """提取图片路径"""
    # 匹配图片路径模式
    img_pattern = r'/tfl/pictures/[^\s\)"]+\.(?:png|jpg|jpeg|gif|bmp)'
    paths = re.findall(img_pattern, content)
    return paths

def extract_docx_content(docx_path: str) -> str:
    """提取 docx 文档内容为纯文本（复制自 docx_summarizer.py）"""
    try:
        if not os.path.exists(docx_path):
            return f"文档文件不存在: {docx_path}"
        
        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 提取表格内容
        table_content = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_content.append(" | ".join(row_data))
        
        # 合并内容
        all_content = []
        if paragraphs:
            all_content.extend(paragraphs)
        if table_content:
            all_content.append("表格内容:")
            all_content.extend(table_content)
        
        return "\n".join(all_content) if all_content else "文档无有效内容"
    except Exception as e:
        return f"文档解析失败: {str(e)}"

@mcp.tool()
async def preprocess_description_field(
    data_file_path: str = "local_data/msg_from_fetcher.json",
    output_file_path: str = "local_data/preprocessed_data.json",
    use_api: bool = True,
    process_documents: bool = False,
    process_images: bool = False
) -> str:
    """
    预处理 TAPD 数据的 description 字段
    
    参数：
        data_file_path (str): 输入数据文件路径
        output_file_path (str): 输出文件路径
        use_api (bool): 是否使用 DeepSeek API 进行内容复述
        process_documents (bool): 是否处理腾讯文档链接（预留功能）
        process_images (bool): 是否处理图片内容（预留功能）
    
    返回：
        str: 处理结果的 JSON 字符串
    """
    try:
        # 读取数据文件
        file_manager = get_file_manager()
        if not os.path.exists(data_file_path):
            return json.dumps({"status": "error", "message": f"数据文件不存在: {data_file_path}"}, ensure_ascii=False)
        
        data = file_manager.load_tapd_data(data_file_path)
        
        processed_count = 0
        api_call_count = 0
        error_count = 0
        results = {"stories": [], "bugs": []}
        
        # 初始化 HTTP 会话
        async with aiohttp.ClientSession() as session:
            # 处理需求数据
            if 'stories' in data:
                for story in data['stories']:
                    try:
                        original_desc = story.get('description', '')
                        if original_desc:
                            # 1. 清理 HTML 样式
                            cleaned_content = clean_html_styles(original_desc)
                            
                            # 2. 提取文档链接和图片路径
                            doc_links = extract_document_links(original_desc)
                            img_paths = extract_image_paths(original_desc)
                            
                            # 3. 使用 API 进行内容复述
                            if use_api and cleaned_content.strip():
                                try:
                                    processed_content = await call_deepseek_api(cleaned_content, session)
                                    api_call_count += 1
                                except Exception as e:
                                    processed_content = cleaned_content
                                    error_count += 1
                                    print(f"API调用失败，使用清理后的内容: {str(e)}")
                            else:
                                processed_content = cleaned_content
                            
                            # 4. 处理文档内容（预留功能）
                            if process_documents and doc_links:
                                for link in doc_links:
                                    processed_content += f"\n\n腾讯文档链接: {link}"
                                    # TODO: 实现文档下载和内容提取
                            
                            # 5. 处理图片内容（预留功能）
                            if process_images and img_paths:
                                for img_path in img_paths:
                                    processed_content += f"\n\n图片路径: {img_path}"
                                    # TODO: 实现图片 OCR 识别
                            
                            # 更新 description 字段
                            story['description'] = processed_content
                            processed_count += 1
                        
                        results['stories'].append(story)
                    except Exception as e:
                        print(f"处理需求 {story.get('id', 'unknown')} 时出错: {str(e)}")
                        error_count += 1
                        results['stories'].append(story)
            
            # 处理缺陷数据
            if 'bugs' in data:
                for bug in data['bugs']:
                    try:
                        original_desc = bug.get('description', '')
                        if original_desc:
                            # 与需求处理逻辑相同
                            cleaned_content = clean_html_styles(original_desc)
                            doc_links = extract_document_links(original_desc)
                            img_paths = extract_image_paths(original_desc)
                            
                            if use_api and cleaned_content.strip():
                                try:
                                    processed_content = await call_deepseek_api(cleaned_content, session)
                                    api_call_count += 1
                                except Exception as e:
                                    processed_content = cleaned_content
                                    error_count += 1
                                    print(f"API调用失败，使用清理后的内容: {str(e)}")
                            else:
                                processed_content = cleaned_content
                            
                            if process_documents and doc_links:
                                for link in doc_links:
                                    processed_content += f"\n\n腾讯文档链接: {link}"
                            
                            if process_images and img_paths:
                                for img_path in img_paths:
                                    processed_content += f"\n\n图片路径: {img_path}"
                            
                            bug['description'] = processed_content
                            processed_count += 1
                        
                        results['bugs'].append(bug)
                    except Exception as e:
                        print(f"处理缺陷 {bug.get('id', 'unknown')} 时出错: {str(e)}")
                        error_count += 1
                        results['bugs'].append(bug)
        
        # 保存处理后的数据
        file_manager.save_json_data(results, output_file_path)
        
        # 返回处理结果
        result_summary = {
            "status": "success",
            "message": "数据预处理完成",
            "statistics": {
                "processed_items": processed_count,
                "api_calls": api_call_count,
                "errors": error_count,
                "output_file": output_file_path
            }
        }
        
        return json.dumps(result_summary, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({"status": "error", "message": f"数据预处理失败: {str(e)}"}, ensure_ascii=False)

@mcp.tool()
def preview_description_cleaning(
    data_file_path: str = "local_data/msg_from_fetcher.json",
    item_count: int = 3
) -> str:
    """
    预览 description 字段清理效果，不实际修改数据
    
    参数：
        data_file_path (str): 数据文件路径
        item_count (int): 预览的条目数量
    
    返回：
        str: 预览结果的 JSON 字符串
    """
    try:
        file_manager = get_file_manager()
        if not os.path.exists(data_file_path):
            return json.dumps({"status": "error", "message": f"数据文件不存在: {data_file_path}"}, ensure_ascii=False)
        
        data = file_manager.load_tapd_data(data_file_path)
        
        preview_results = []
        count = 0
        
        # 预览需求数据
        if 'stories' in data and count < item_count:
            for story in data['stories']:
                if count >= item_count:
                    break
                
                original_desc = story.get('description', '')
                if original_desc:
                    cleaned_desc = clean_html_styles(original_desc)
                    doc_links = extract_document_links(original_desc)
                    img_paths = extract_image_paths(original_desc)
                    
                    preview_results.append({
                        "id": story.get('id', 'unknown'),
                        "type": "story",
                        "name": story.get('name', ''),
                        "original_length": len(original_desc),
                        "cleaned_length": len(cleaned_desc),
                        "original_preview": original_desc[:200] + "..." if len(original_desc) > 200 else original_desc,
                        "cleaned_preview": cleaned_desc[:200] + "..." if len(cleaned_desc) > 200 else cleaned_desc,
                        "document_links": doc_links,
                        "image_paths": img_paths
                    })
                    count += 1
        
        # 预览缺陷数据
        if 'bugs' in data and count < item_count:
            for bug in data['bugs']:
                if count >= item_count:
                    break
                
                original_desc = bug.get('description', '')
                if original_desc:
                    cleaned_desc = clean_html_styles(original_desc)
                    doc_links = extract_document_links(original_desc)
                    img_paths = extract_image_paths(original_desc)
                    
                    preview_results.append({
                        "id": bug.get('id', 'unknown'),
                        "type": "bug",
                        "title": bug.get('title', ''),
                        "original_length": len(original_desc),
                        "cleaned_length": len(cleaned_desc),
                        "original_preview": original_desc[:200] + "..." if len(original_desc) > 200 else original_desc,
                        "cleaned_preview": cleaned_desc[:200] + "..." if len(cleaned_desc) > 200 else cleaned_desc,
                        "document_links": doc_links,
                        "image_paths": img_paths
                    })
                    count += 1
        
        return json.dumps({
            "status": "success",
            "preview_count": len(preview_results),
            "results": preview_results
        }, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({"status": "error", "message": f"预览失败: {str(e)}"}, ensure_ascii=False)

if __name__ == "__main__":
    # 测试代码
    async def test():
        result = await preprocess_description_field(
            data_file_path="local_data/msg_from_fetcher.json",
            output_file_path="local_data/test_preprocessed.json",
            use_api=False  # 测试时不使用API
        )
        print(result)
    
    # 运行测试
    # asyncio.run(test())
    
    # 预览清理效果
    preview_result = preview_description_cleaning()
    print("预览结果:")
    print(preview_result)
