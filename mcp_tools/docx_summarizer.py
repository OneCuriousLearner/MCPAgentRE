"""
docx_summarizer.py

功能：
- 输入 .docx 文档路径，自动提取正文内容。
- 生成简要摘要（如截取前N段）。
- 作为 MCP 工具注册，便于 AI 自动调用。

依赖：python-docx
"""
from typing import Optional
from docx import Document
import os
from mcp.server.fastmcp import FastMCP
import json
import shutil
import uuid
import csv

mcp = FastMCP("docx")

@mcp.tool()
def summarize_docx(docx_path: str, max_paragraphs: int = 5) -> str:
    """
    读取 docx 文档并生成摘要和完整内容的 JSON 数据，同时提取图片和表格
    
    参数：
        docx_path (str): .docx 文件路径
        max_paragraphs (int): 摘要最多包含的段落数，默认5
    返回：
        str: JSON 字符串，包含所有段落、摘要、图片和表格信息
    """
    if not os.path.exists(docx_path):
        return json.dumps({"status": "error", "message": f"文件不存在: {docx_path}"}, ensure_ascii=False)
    try:
        doc = Document(docx_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        para_list = [
            {"index": i+1, "text": text}
            for i, text in enumerate(paragraphs)
        ]
        summary = "\n".join(paragraphs[:max_paragraphs]) if paragraphs else "文档无有效正文内容。"

        # 1. 提取图片
        pictures_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../pictures_data'))
        os.makedirs(pictures_dir, exist_ok=True)
        picture_files = []
        rels = doc.part.rels
        for rel in rels:
            rel_obj = rels[rel]
            if rel_obj.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                img = rel_obj.target_part
                img_data = img.blob
                ext = os.path.splitext(img.partname)[1]
                img_name = f"{os.path.splitext(os.path.basename(docx_path))[0]}_{uuid.uuid4().hex[:8]}{ext}"
                img_path = os.path.join(pictures_dir, img_name)
                with open(img_path, 'wb') as f:
                    f.write(img_data)
                picture_files.append(img_name)

        # 2. 提取表格为 CSV
        tables_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../excel_data'))
        os.makedirs(tables_dir, exist_ok=True)
        table_files = []
        for idx, table in enumerate(doc.tables):
            csv_name = f"{os.path.splitext(os.path.basename(docx_path))[0]}_table_{idx+1}.csv"
            csv_path = os.path.join(tables_dir, csv_name)
            with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                for row in table.rows:
                    writer.writerow([cell.text.strip() for cell in row.cells])
            table_files.append(csv_name)

        result = {
            "status": "success",
            "summary": summary,
            "paragraphs": para_list,
            "total_paragraphs": len(para_list),
            "pictures": picture_files,
            "tables": table_files
        }
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"status": "error", "message": f"解析文档失败: {str(e)}"}, ensure_ascii=False)

if __name__ == "__main__":
    # 示例用法
    test_path = "../docx_sum/2.docx"
    result_json = summarize_docx(test_path)
    print(result_json)
    # 输出到主文件夹
    with open("../docx_summary.json", "w", encoding="utf-8") as f:
        f.write(result_json)
    print("已将文本摘要输出到主文件夹，图片信息输出到pictures_data文件夹，表格信息输出到excel_data文件夹。") 